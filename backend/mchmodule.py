from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


dic = {'От кого':'payer', 'БИК':'bik', 'Номер счета':'num_account', 'За что':'nds', 'Сумма к оплате':'sum_pay'}
arr = ['От кого', 'БИК', 'Номер счета', 'За что', 'Сумма к оплате']

def make_charge(num, data, filename):
    document = Document()
    document.add_heading('ПЛАТЁЖ №' + str(num))
    document.add_paragraph("")
    table = document.add_table(rows=5, cols=2)
    table.style = 'TableGrid'
    for i in range(5):
        row = table.rows[i]
        row.cells[0].text = arr[i]
        row.cells[1].text = data[dic[arr[i]]]
    document.add_paragraph("")
    p = document.add_paragraph("_________________________________")
    r = p.add_run()
    r.add_picture('mark.png')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.save(filename)
    print('save')